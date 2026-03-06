"""Spatial Attention Detection 참고 문헌 및 Object Detection 근거 분석 Word 문서 생성"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── 한글 폰트 설정 ──
# 우선순위: 맑은 고딕(Windows) > Apple SD Gothic Neo(Mac) > Noto Sans KR(Linux/Web)
FONT_NAME = "맑은 고딕"
FONT_FALLBACK = "Malgun Gothic"


def set_korean_font(run, font_name=FONT_NAME, fallback=FONT_FALLBACK):
    """Latin + East Asian 폰트를 모두 설정하여 한글 깨짐 방지"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), fallback)
    rFonts.set(qn("w:hAnsi"), fallback)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), fallback)


def set_style_korean_font(style, font_name=FONT_NAME, fallback=FONT_FALLBACK):
    """스타일의 rPr에 한글 폰트 설정"""
    style.font.name = fallback
    # East Asian 폰트 직접 설정
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), fallback)
    rFonts.set(qn("w:hAnsi"), fallback)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), fallback)


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="2E4057"):
    """스타일이 적용된 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 헤더 행
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_korean_font(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            set_korean_font(run)
            # 짝수 행 배경색
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    # 열 너비 설정
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def create_document():
    doc = Document()

    # ── 스타일 설정 (Latin + East Asian 모두) ──
    style = doc.styles["Normal"]
    set_style_korean_font(style)
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Heading 스타일
    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        set_style_korean_font(h_style)
        h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ════════════════════════════════════════════════════════════
    #  표지
    # ════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Spatial Attention Detection 모델\n참고 문헌 및 학술적 근거 분석")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    set_korean_font(run)

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("FBM 결함 분류 프로젝트 — Eval 4 모델 분석")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    set_korean_font(run)

    doc.add_paragraph("")
    doc.add_paragraph("")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"작성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    set_korean_font(run)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    #  목차
    # ════════════════════════════════════════════════════════════
    doc.add_heading("목차", level=1)
    toc_items = [
        "1. 개요",
        "2. Eval 4 모델 구조 요약",
        "3. 참고 문헌",
        "   3.1 Class Activation Mapping (CAM)",
        "   3.2 Spatial Attention 메커니즘",
        "   3.3 Weakly-Supervised Detection/Localization",
        "   3.4 Multi-label Classification with Attention",
        "   3.5 결함 검출 및 Survey",
        "4. Object Detection 해당 여부 분석",
        "   4.1 학술적 정의 비교",
        "   4.2 Detection이라 부를 수 있는 근거",
        "   4.3 Detection이라 부르기 어려운 이유",
        "   4.4 권장 용어",
        "5. 결론",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    #  1. 개요
    # ════════════════════════════════════════════════════════════
    doc.add_heading("1. 개요", level=1)
    doc.add_paragraph(
        "본 문서는 FBM(Flat Board Module) 결함 분류 프로젝트의 Eval 4에서 사용된 "
        "Spatial Attention Detection 모델의 학술적 근거를 정리한 것입니다. "
        "해당 모델은 이미지 레벨 라벨(image-level label)만을 사용하여 학습하면서도, "
        "클래스별 독립적인 공간 주의 맵(Spatial Attention Map)을 생성하여 "
        "각 결함의 위치를 시각적으로 확인할 수 있는 구조입니다."
    )
    doc.add_paragraph(
        "본 문서에서는 (1) 모델 구조와 관련된 핵심 참고 문헌을 분류별로 정리하고, "
        "(2) 이 모델을 'Object Detection'이라고 볼 수 있는지에 대한 학술적 근거를 분석합니다."
    )

    # ════════════════════════════════════════════════════════════
    #  2. 모델 구조 요약
    # ════════════════════════════════════════════════════════════
    doc.add_heading("2. Eval 4 모델 구조 요약", level=1)

    doc.add_paragraph(
        "Eval 4 모델은 표준 CNN Classifier(Eval 1~3)와 달리, Global Average Pooling을 "
        "사용하지 않고 클래스별 독립적인 Spatial Attention Head를 통해 "
        "각 결함 유형의 공간 위치를 보존하며 탐지합니다."
    )

    # 구조 비교 테이블
    add_styled_table(doc,
        ["구분", "Standard CNN (Eval 1~3)", "Spatial Attention Model (Eval 4)"],
        [
            ["공간 정보 처리", "AdaptiveAvgPool → 공간 정보 완전 압축", "클래스별 Attention Map → 공간 위치 보존"],
            ["분류 방식", "단일 FC → 전체 클래스 동시 출력", "클래스별 독립 FC → 개별 판단"],
            ["중첩 대응", "특징 혼합으로 중첩 시 혼동", "공간 분리로 독립 탐지"],
            ["파라미터 수", "~406K", "~620K"],
            ["해석 가능성", "Post-hoc (Grad-CAM 필요)", "Built-in (Attention Map 직접 출력)"],
        ],
        col_widths=[3, 6.5, 6.5]
    )

    doc.add_paragraph("")

    arch_p = doc.add_paragraph()
    arch_run = arch_p.add_run(
        "핵심 구조: Backbone(Conv×4) → Per-Class Attention Head(Conv1×1 bottleneck + Sigmoid) "
        "→ Attended Feature Aggregation(feature × attention → spatial mean) → Per-Class FC → Logits"
    )
    arch_run.font.size = Pt(9)
    arch_run.italic = True
    set_korean_font(arch_run)

    # ════════════════════════════════════════════════════════════
    #  3. 참고 문헌
    # ════════════════════════════════════════════════════════════
    doc.add_heading("3. 참고 문헌", level=1)

    # ── 3.1 CAM ──
    doc.add_heading("3.1 Class Activation Mapping (CAM)", level=2)
    doc.add_paragraph(
        "CAM 계열은 CNN이 분류를 수행할 때 \"어디를 보고 판단하는지\"를 시각화하는 기법입니다. "
        "Eval 4 모델의 Attention Map은 CAM의 학습 가능한(learnable) 확장판에 해당합니다."
    )

    refs_cam = [
        [
            "[1]",
            "Learning Deep Features for Discriminative Localization",
            "Zhou, Khosla, Lapedriza, Oliva, Torralba",
            "CVPR 2016",
            "CAM 원조 논문. GAP + FC weights로 클래스별 spatial heatmap 생성. "
            "Eval4의 per-class attention head는 이 개념의 학습 가능한 확장판으로, "
            "CAM이 선형 가중합인 반면 Eval4는 Conv bottleneck + Sigmoid로 더 높은 표현력을 가짐."
        ],
        [
            "[2]",
            "Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization",
            "Selvaraju, Cogswell, Das, Vedantam, Parikh, Batra",
            "ICCV 2017",
            "Gradient 기반으로 임의 CNN에서 클래스별 spatial map 생성. "
            "Eval4의 attention map이 Grad-CAM과 동일한 역할을 하되, "
            "post-hoc이 아닌 forward pass에 내장(baked-in)된 차이점이 있음."
        ],
    ]
    add_styled_table(doc,
        ["#", "논문 제목", "저자", "학회", "Eval 4와의 관련성"],
        refs_cam,
        col_widths=[1, 4, 3, 1.5, 6.5]
    )

    # ── 3.2 Spatial Attention ──
    doc.add_heading("3.2 Spatial Attention 메커니즘", level=2)
    doc.add_paragraph(
        "CNN에 attention 메커니즘을 도입하여 \"어디에 집중할지\"를 학습하는 기법들입니다. "
        "Eval 4의 attention head 구조는 이 계열의 연구에 기반합니다."
    )

    refs_attn = [
        [
            "[3]",
            "Squeeze-and-Excitation Networks (SE-Net)",
            "Hu, Shen, Sun",
            "CVPR 2018",
            "채널 attention의 대표 논문 (ILSVRC 2017 우승). "
            "Bottleneck(256→64→1) + Sigmoid 패턴이 Eval4의 attention head와 동일하나, "
            "SE-Net은 채널 차원, Eval4는 공간(spatial) 차원에 적용."
        ],
        [
            "[4]",
            "CBAM: Convolutional Block Attention Module",
            "Woo, Park, Lee, Kweon",
            "ECCV 2018",
            "채널 + 공간 attention 결합. Eval4와 가장 구조적으로 유사 — "
            "CBAM의 spatial branch가 Conv로 spatial weight를 생성하는 것과 동일. "
            "차이점: CBAM은 하나의 공유 spatial map, Eval4는 클래스별 독립 spatial map."
        ],
        [
            "[5]",
            "Residual Attention Network for Image Classification",
            "Wang, Jiang, Yang, Wu, Qian, Li",
            "CVPR 2017",
            "Soft attention(feature × attention_map) 패턴의 근거. "
            "Eval4의 f * attn 연산이 이 패턴을 따름."
        ],
    ]
    add_styled_table(doc,
        ["#", "논문 제목", "저자", "학회", "Eval 4와의 관련성"],
        refs_attn,
        col_widths=[1, 4, 3, 1.5, 6.5]
    )

    # ── 3.3 WSOD/WSOL ──
    doc.add_heading("3.3 Weakly-Supervised Detection / Localization", level=2)
    doc.add_paragraph(
        "이미지 레벨 라벨만으로 객체의 공간 위치를 학습하는 기법입니다. "
        "Eval 4 모델은 BBox 라벨 없이 BCEWithLogitsLoss만으로 공간 주의를 학습하므로, "
        "이 범주에 해당합니다."
    )

    refs_wsod = [
        [
            "[6]",
            "Weakly Supervised Deep Detection Networks (WSDDN)",
            "Bilen, Vedaldi",
            "CVPR 2016",
            "WSOD 핵심 논문. 이미지 레벨 라벨만으로 classification loss를 통해 "
            "공간 선택(spatial selection) 학습 — Eval4의 핵심 원리와 동일. "
            "WSDDN은 softmax over regions, Eval4는 sigmoid attention maps 사용."
        ],
        [
            "[7]",
            "Learning Spatial Regularization with Image-Level Supervisions "
            "for Multi-label Image Classification",
            "Zhu, Li, Ouyang, Yu, Wang",
            "CVPR 2017",
            "Eval4와 가장 직접적으로 관련된 논문. "
            "이미지 레벨 라벨로 per-label attention map을 생성하여 multi-label 분류 — "
            "구조적으로 거의 동일한 접근."
        ],
    ]
    add_styled_table(doc,
        ["#", "논문 제목", "저자", "학회", "Eval 4와의 관련성"],
        refs_wsod,
        col_widths=[1, 4, 3, 1.5, 6.5]
    )

    # ── 3.4 Multi-label ──
    doc.add_heading("3.4 Multi-label Classification with Per-Class Attention", level=2)
    doc.add_paragraph(
        "다중 라벨 분류에서 클래스별 독립적인 attention을 사용하는 연구들입니다."
    )

    refs_ml = [
        [
            "[8]",
            "Multi-label Image Recognition by Recurrently Discovering Attentional Regions",
            "Wang, Chen, Li, Xu, Lin",
            "ICCV 2017",
            "클래스별 공간 영역을 순차 탐색하여 multi-label 인식. "
            "Eval4는 순차적(recurrent) 대신 병렬(parallel) attention head 구조로 동일 목적 달성."
        ],
    ]
    add_styled_table(doc,
        ["#", "논문 제목", "저자", "학회", "Eval 4와의 관련성"],
        refs_ml,
        col_widths=[1, 4, 3, 1.5, 6.5]
    )

    # ── 3.5 Survey ──
    doc.add_heading("3.5 결함 검출 및 Survey", level=2)

    refs_survey = [
        [
            "[9]",
            "Weakly Supervised Object Localization and Detection: A Survey",
            "Zhang et al.",
            "IEEE TPAMI 2022",
            "WSOL/WSOD 기법 종합 서베이. Eval4의 학술적 위치를 정의하는 데 인용 가능. "
            "WSOL과 WSOD의 엄밀한 정의를 제공."
        ],
    ]
    add_styled_table(doc,
        ["#", "논문 제목", "저자", "학회", "Eval 4와의 관련성"],
        refs_survey,
        col_widths=[1, 4, 3, 1.5, 6.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    #  4. Object Detection 해당 여부 분석
    # ════════════════════════════════════════════════════════════
    doc.add_heading("4. Object Detection 해당 여부 분석", level=1)

    # ── 4.1 정의 비교 ──
    doc.add_heading("4.1 학술적 정의 비교", level=2)
    doc.add_paragraph(
        "Zhang et al. (IEEE TPAMI, 2022)의 서베이 논문에 따른 학술적 정의와 "
        "Eval 4 모델의 해당 여부를 비교합니다."
    )

    add_styled_table(doc,
        ["구분", "정의", "학습 라벨", "출력", "Eval 4 해당"],
        [
            ["Object Detection\n(정통)", "BBox regression +\nClassification", "BBox 좌표 +\n클래스", "(x,y,w,h) +\nclass", "해당 없음"],
            ["WSOL\n(Weakly-Supervised\nObject Localization)", "이미지 라벨로\n단일 객체 위치 학습", "이미지 레벨\n라벨", "Heatmap /\nBBox", "부분 해당"],
            ["WSOD\n(Weakly-Supervised\nObject Detection)", "이미지 라벨로\n다중 객체 탐지 학습", "이미지 레벨\n라벨", "BBox +\nclass", "부분 해당"],
            ["Attention-based\nMulti-label\nClassification", "클래스별 attention으로\nmulti-label 분류", "이미지 레벨\n라벨", "Heatmap +\nclass", "해당"],
        ],
        col_widths=[3.5, 3.5, 2.5, 2.5, 2.5]
    )

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run(
        "결론: Eval 4 모델은 \"Attention-based Multi-label Classification with "
        "Implicit Spatial Localization\" 범주에 가장 정확하게 해당하며, "
        "WSOL/WSOD의 특성을 부분적으로 공유합니다."
    )
    run.bold = True
    set_korean_font(run)

    # ── 4.2 Detection이라 부를 수 있는 근거 ──
    doc.add_heading("4.2 Detection이라 부를 수 있는 근거", level=2)

    reasons_for = [
        (
            "공간 위치 식별",
            "Eval4 모델은 각 결함이 \"어디에\" 있는지를 attention map으로 출력합니다. "
            "이것은 detection의 핵심 요소(localization)에 해당합니다."
        ),
        (
            "클래스별 독립 탐지",
            "7개 클래스 각각에 대해 독립적인 spatial map을 생성하여 서로 다른 영역을 탐지합니다. "
            "이는 multi-object detection과 동일한 목적입니다."
        ),
        (
            "중첩 불량 분리",
            "하나의 이미지에서 여러 결함을 공간적으로 분리하여 동시 탐지합니다. "
            "이것이 바로 detection이 필요한 이유이자, 표준 CNN 대비 Eval4가 우수한 이유입니다."
        ),
        (
            "학술적 관례",
            "WSDDN(Bilen et al., 2016)처럼 이미지 레벨 라벨만으로 공간 학습을 수행하는 모델을 "
            "학술적으로 \"weakly-supervised detection\"이라고 널리 칭합니다."
        ),
        (
            "해석 가능한 출력",
            "일반 CNN은 \"있다/없다\"만 출력하지만, Eval4는 \"어디에 있다\"까지 출력합니다. "
            "이는 분류(classification)를 넘어선 탐지(detection)의 정의에 부합합니다."
        ),
    ]

    for i, (title, desc) in enumerate(reasons_for, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"근거 {i}: {title}")
        run_num.bold = True
        run_num.font.size = Pt(10)
        set_korean_font(run_num)
        doc.add_paragraph(desc)

    # ── 4.3 Detection이라 부르기 어려운 이유 ──
    doc.add_heading("4.3 Detection이라 부르기 어려운 이유", level=2)

    reasons_against = [
        (
            "BBox 미출력",
            "정통 Object Detection은 (x, y, w, h) 좌표를 예측하지만, "
            "Eval4는 연속적 heatmap만 출력합니다. "
            "학계에서 \"detection\"은 통상 bbox 수준의 예측을 의미합니다."
        ),
        (
            "인스턴스 구분 불가",
            "같은 클래스의 여러 인스턴스가 존재할 경우, 이를 개별적으로 구분하여 탐지할 수 없습니다."
        ),
        (
            "Localization 정확도에 대한 직접적 supervision 부재",
            "Ground truth가 이미지 단위 이진 라벨이므로, attention map이 정확한 위치를 "
            "가리키도록 강제하는 직접적 학습 신호가 없습니다."
        ),
    ]

    for i, (title, desc) in enumerate(reasons_against, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f"한계 {i}: {title}")
        run_num.bold = True
        run_num.font.size = Pt(10)
        set_korean_font(run_num)
        doc.add_paragraph(desc)

    # ── 4.4 권장 용어 ──
    doc.add_heading("4.4 권장 용어", level=2)

    add_styled_table(doc,
        ["사용 맥락", "권장 표현", "비고"],
        [
            [
                "논문 / 공식 보고서",
                "Weakly-supervised multi-label classification\nwith learned spatial attention",
                "가장 학술적으로 정확한 표현"
            ],
            [
                "기술 보고서 (간결)",
                "Attention-based weakly-supervised\ndefect localization",
                "Localization 능력을 강조"
            ],
            [
                "내부 / 구두 소통",
                "Detection 모델",
                "실용적으로 문제없음.\n단, 정통 Object Detection과\n구별 필요 시 주석 추가"
            ],
        ],
        col_widths=[3.5, 6, 5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    #  5. 결론
    # ════════════════════════════════════════════════════════════
    doc.add_heading("5. 결론", level=1)

    doc.add_paragraph(
        "Eval 4의 Spatial Attention Detection 모델은 학술적으로 다음과 같이 위치합니다:"
    )

    conclusions = [
        "CAM(Zhou et al., 2016)의 학습 가능한 확장판으로, 클래스별 독립적인 spatial attention을 "
        "forward pass에 내장하여 post-hoc 시각화 없이도 모델의 판단 근거를 제공합니다.",

        "CBAM(Woo et al., 2018)의 spatial attention branch와 구조적으로 유사하나, "
        "클래스별 독립 attention head를 사용하여 multi-label 환경에서 공간 분리를 달성합니다.",

        "WSDDN(Bilen et al., 2016) 및 Zhu et al.(2017)과 동일한 원리로, "
        "이미지 레벨 라벨만으로 공간적 localization을 학습하는 weakly-supervised 방식입니다.",

        "엄밀한 의미의 Object Detection(bbox regression)은 아니지만, "
        "공간 위치를 식별하고 클래스별 독립 탐지를 수행한다는 점에서 "
        "\"Weakly-Supervised Detection\"이라는 명칭이 학술적으로 타당합니다.",
    ]

    for i, text in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        set_korean_font(run)
        run2 = p.add_run(text)
        set_korean_font(run2)

    doc.add_paragraph("")

    final = doc.add_paragraph()
    run = final.add_run(
        "향후 BBox 라벨을 추가하여 학습할 경우, 정통 Object Detection으로의 확장이 가능하며, "
        "이 경우 더 정확한 공간 localization과 인스턴스 수준의 탐지가 기대됩니다."
    )
    run.italic = True
    set_korean_font(run)

    # ════════════════════════════════════════════════════════════
    #  참고 문헌 목록 (간결 버전)
    # ════════════════════════════════════════════════════════════
    doc.add_heading("참고 문헌", level=1)

    references = [
        "[1] Zhou, B., Khosla, A., Lapedriza, A., Oliva, A., & Torralba, A. (2016). "
        "\"Learning Deep Features for Discriminative Localization.\" CVPR 2016.",

        "[2] Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). "
        "\"Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization.\" ICCV 2017.",

        "[3] Hu, J., Shen, L., & Sun, G. (2018). "
        "\"Squeeze-and-Excitation Networks.\" CVPR 2018.",

        "[4] Woo, S., Park, J., Lee, J.-Y., & Kweon, I. S. (2018). "
        "\"CBAM: Convolutional Block Attention Module.\" ECCV 2018.",

        "[5] Wang, F., Jiang, M., Qian, C., Yang, S., Li, C., Zhang, H., Wang, X., & Tang, X. (2017). "
        "\"Residual Attention Network for Image Classification.\" CVPR 2017.",

        "[6] Bilen, H. & Vedaldi, A. (2016). "
        "\"Weakly Supervised Deep Detection Networks.\" CVPR 2016.",

        "[7] Zhu, F., Li, H., Ouyang, W., Yu, N., & Wang, X. (2017). "
        "\"Learning Spatial Regularization with Image-Level Supervisions "
        "for Multi-label Image Classification.\" CVPR 2017.",

        "[8] Wang, Z., Chen, T., Li, G., Xu, R., & Lin, L. (2017). "
        "\"Multi-label Image Recognition by Recurrently Discovering Attentional Regions.\" ICCV 2017.",

        "[9] Zhang, D. et al. (2022). "
        "\"Weakly Supervised Object Localization and Detection: A Survey.\" IEEE TPAMI 2022.",
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        for run in p.runs:
            run.font.size = Pt(9)
            set_korean_font(run)

    # ── 저장 ──
    output_path = "/home/user/fbm_classification/docs/spatial_attention_references.docx"
    doc.save(output_path)
    print(f"문서 생성 완료: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
